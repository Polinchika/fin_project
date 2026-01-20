from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt

def replace_placeholders(doc: Document, values: dict):
    for paragraph in doc.paragraphs:
        original_text = paragraph.text

        new_text = original_text
        for key, val in values.items():
            placeholder = f"{{{{{key}}}}}"
            new_text = new_text.replace(placeholder, val)

        if new_text != original_text:
            paragraph.clear()
            paragraph.add_run(new_text)


def generate_pdf(text: str, title: str) -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    y = height - 40

    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, title)
    y -= 30

    c.setFont("Helvetica", 10)

    for line in text.splitlines():
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 40
        c.drawString(40, y, line)
        y -= 14

    c.save()
    buffer.seek(0)
    return buffer


def generate_docx(values: dict) -> BytesIO:
    doc = Document("template_1.docx")

    replace_placeholders(doc, values)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
